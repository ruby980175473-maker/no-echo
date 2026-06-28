# ============================================================
# NO ECHO · DOCX 解析器
# Sprint 2 核心实现
#
# 设计原则：
# - 每个字段由独立函数解析（parse_title / parse_headings / ...）
# - 解析失败返回 None / []，不抛出异常，不中断整体流程
# - 所有函数入参为 doc.paragraphs（已解析好的段落列表），
#   避免重复打开文件
# ============================================================

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from app.modules.parser.models import (
    DocumentMetadata,
    DocumentModel,
    HeadingItem,
    ImageItem,
    ParagraphItem,
    TableItem,
)

# ── 样式名 → 标题层级映射（中英文 Word 样式名兼容）──────────────
_HEADING_STYLE_LEVEL: dict[str, int] = {
    "heading 1": 1, "标题 1": 1, "标题1": 1,
    "heading 2": 2, "标题 2": 2, "标题2": 2,
    "heading 3": 3, "标题 3": 3, "标题3": 3,
}

# ── 节标记关键词 ───────────────────────────────────────────────
_ABSTRACT_CN_MARKERS  = {"摘要", "中文摘要", "内容摘要"}
_ABSTRACT_EN_MARKERS  = {"abstract", "english abstract"}
_KEYWORDS_CN_PREFIXES = ("关键词", "关键字")
_KEYWORDS_EN_PREFIXES = ("keywords", "key words")
_REFERENCE_MARKERS    = {"参考文献", "references", "bibliography", "[参考文献]"}

# 参考文献条目：[1] 或 1. 开头
_REF_ITEM_RE = re.compile(r'^\s*[\[【]?\d+[\]】]?[\.\s、]')

# 数字编号标题：1. 或 1.1 或 1.1.1（后面有空格或中文）
_NUMBERED_HEADING_RE = re.compile(r'^(\d+)(\.\d+)*[\s\u3000\u4e00-\u9fff]')


# ── 内部工具函数 ────────────────────────────────────────────────

def _heading_level(para) -> Optional[int]:
    """
    获取段落的标题层级（1/2/3），非标题返回 None。
    检测顺序：① Word 样式名 → ② 数字编号模式
    """
    style = (para.style.name or "").lower().strip()
    if style in _HEADING_STYLE_LEVEL:
        return _HEADING_STYLE_LEVEL[style]

    text = para.text.strip()
    if text and _NUMBERED_HEADING_RE.match(text):
        dots = text.split()[0].count('.') if text.split() else 0
        return min(dots + 1, 3)

    return None


def _is_marker(text: str, markers: set[str]) -> bool:
    """完全匹配（忽略大小写）"""
    return text.strip().lower() in {m.lower() for m in markers}


def _starts_with(text: str, prefixes: tuple) -> bool:
    """前缀匹配（忽略大小写）"""
    t = text.strip().lower()
    return any(t.startswith(p.lower()) for p in prefixes)


def _split_keywords(raw: str, separators: str = r'[；;，,、\s]+') -> list[str]:
    """分割关键词字符串，过滤空项"""
    return [k.strip() for k in re.split(separators, raw) if k.strip()]


def _count_words(text: str) -> int:
    """
    统计词数：CJK 字符各计 1 词，英文按空格分词
    """
    cjk = len(re.findall(r'[\u4e00-\u9fff\u3400-\u4dbf]', text))
    en  = len(re.findall(r'[a-zA-Z]+', text))
    return cjk + en


# ── 独立解析函数 ────────────────────────────────────────────────

def parse_title(paras: list, headings: list[HeadingItem]) -> Optional[str]:
    """
    提取论文标题。
    策略（依次尝试）：
    1. 样式名含 "title" / "题目" 的段落
    2. 第一个一级标题
    3. 文档第一个非空段落（兜底，至少 4 个字符）
    """
    for para in paras:
        style = (para.style.name or "").lower()
        if "title" in style or "题目" in style:
            t = para.text.strip()
            if t:
                return t

    for h in headings:
        if h.level == 1:
            return h.text

    for para in paras:
        t = para.text.strip()
        if t and len(t) >= 4:
            return t

    return None


def parse_headings(paras: list) -> list[HeadingItem]:
    """
    提取所有标题，含层级和编号信息。
    """
    result = []
    for i, para in enumerate(paras):
        level = _heading_level(para)
        if level is None:
            continue
        text = para.text.strip()
        if not text:
            continue

        # 尝试提取编号（如 "1.1"）
        numbering: Optional[str] = None
        m = _NUMBERED_HEADING_RE.match(text)
        if m:
            num_part = text[:m.end()].strip()
            numbering = num_part if num_part else None

        result.append(HeadingItem(index=i, level=level, text=text, numbering=numbering))
    return result


def parse_paragraphs(paras: list) -> list[ParagraphItem]:
    """
    提取所有非标题非空段落（含摘要，供 AIGC / 重复风险模块全文扫描）。
    """
    result = []
    for i, para in enumerate(paras):
        text = para.text.strip()
        if not text:
            continue
        if _heading_level(para) is not None:
            continue
        result.append(ParagraphItem(
            index=i,
            text=text,
            style_name=para.style.name or "Unknown",
            char_count=len(text.replace(" ", "")),
        ))
    return result


def parse_abstract_cn(paras: list) -> Optional[str]:
    """
    提取中文摘要正文。
    支持两种格式：
    - 独立行：「摘要」后接内容段落
    - 同行：「摘要：内容...」
    """
    collecting = False
    lines: list[str] = []

    for para in paras:
        text = para.text.strip()
        if not text:
            continue

        if not collecting:
            if _is_marker(text, _ABSTRACT_CN_MARKERS):
                collecting = True
                continue
            # 同行格式
            for prefix in ("摘要：", "摘要:", "中文摘要：", "中文摘要:"):
                if text.startswith(prefix):
                    after = text[len(prefix):].strip()
                    if after:
                        lines.append(after)
                    collecting = True
                    break
        else:
            # 停止条件
            if (_is_marker(text, _ABSTRACT_EN_MARKERS)
                    or _starts_with(text, _KEYWORDS_CN_PREFIXES)
                    or _starts_with(text, _KEYWORDS_EN_PREFIXES)
                    or _is_marker(text, _REFERENCE_MARKERS)
                    or _heading_level(para) is not None):
                break
            lines.append(text)

    return "\n".join(lines) if lines else None


def parse_abstract_en(paras: list) -> Optional[str]:
    """提取英文摘要正文"""
    collecting = False
    lines: list[str] = []

    for para in paras:
        text = para.text.strip()
        if not text:
            continue

        if not collecting:
            if _is_marker(text, _ABSTRACT_EN_MARKERS):
                collecting = True
                continue
            for prefix in ("abstract:", "abstract："):
                if text.lower().startswith(prefix):
                    after = text[len(prefix):].strip()
                    if after:
                        lines.append(after)
                    collecting = True
                    break
        else:
            if (_starts_with(text, _KEYWORDS_EN_PREFIXES)
                    or _starts_with(text, _KEYWORDS_CN_PREFIXES)
                    or _is_marker(text, _REFERENCE_MARKERS)
                    or _heading_level(para) is not None):
                break
            lines.append(text)

    return "\n".join(lines) if lines else None


def parse_keywords_cn(paras: list) -> list[str]:
    """提取中文关键词列表"""
    for para in paras:
        text = para.text.strip()
        if _starts_with(text, _KEYWORDS_CN_PREFIXES):
            for prefix in _KEYWORDS_CN_PREFIXES:
                if text.lower().startswith(prefix.lower()):
                    raw = text[len(prefix):].lstrip("：: ").strip()
                    return _split_keywords(raw)
    return []


def parse_keywords_en(paras: list) -> list[str]:
    """提取英文关键词列表"""
    for para in paras:
        text = para.text.strip()
        if _starts_with(text, _KEYWORDS_EN_PREFIXES):
            for prefix in _KEYWORDS_EN_PREFIXES:
                if text.lower().startswith(prefix.lower()):
                    raw = text[len(prefix):].lstrip("：:; ").strip()
                    # 英文关键词用逗号或分号分隔均支持
                    sep = "," if "," in raw else ";"
                    return [k.strip() for k in raw.split(sep) if k.strip()]
    return []


def parse_references(paras: list) -> list[str]:
    """
    提取参考文献条目列表。
    策略：找到"参考文献"节后，收集以 [数字] 或 数字. 开头的段落。
    多行条目：非条目开头的段落追加到上一条目末尾。
    """
    collecting = False
    result: list[str] = []

    for para in paras:
        text = para.text.strip()
        if not text:
            continue

        if not collecting:
            if _is_marker(text, _REFERENCE_MARKERS):
                collecting = True
            continue

        # 进入参考文献区
        if _REF_ITEM_RE.match(text):
            result.append(text)
        elif result:
            # 如果遇到新的一级标题，停止
            if _heading_level(para) == 1:
                break
            # 否则视为上一条目的续行
            result[-1] += " " + text

    return result


def parse_images(doc, paras: list) -> list[ImageItem]:
    """
    提取图片信息（inline shapes）。
    图题：优先检查图片所在段落的文字；其次检查下一段落。
    """
    result: list[ImageItem] = []
    image_count = 0

    for i, para in enumerate(paras):
        has_image = bool(para._element.findall('.//' + qn('a:graphic')))
        if not has_image:
            continue

        image_count += 1
        caption: Optional[str] = None

        # 当前段落文字可能含图题
        para_text = para.text.strip()
        if para_text and (para_text.startswith("图") or para_text.lower().startswith("figure")):
            caption = para_text
        elif i + 1 < len(paras):
            next_text = paras[i + 1].text.strip()
            if next_text and (next_text.startswith("图") or next_text.lower().startswith("figure")):
                caption = next_text

        result.append(ImageItem(index=image_count, caption=caption, paragraph_index=i))

    return result


def parse_tables(doc) -> list[TableItem]:
    """提取表格信息（行数、列数）"""
    result: list[TableItem] = []
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns) if table.rows else 0
        result.append(TableItem(
            index=i + 1,
            caption=None,       # 表题检测复杂，Sprint 3 实现
            rows=rows,
            cols=cols,
            paragraph_index=-1, # 精确定位需遍历 XML，Sprint 3 实现
        ))
    return result


def parse_metadata(
    file_path: Path,
    paras: list,
    headings: list[HeadingItem],
    images: list[ImageItem],
    tables: list[TableItem],
    references: list[str],
    abstract_cn: Optional[str],
    abstract_en: Optional[str],
    keywords_cn: list[str],
    keywords_en: list[str],
) -> DocumentMetadata:
    """计算文档元数据"""
    all_text = " ".join(p.text.strip() for p in paras if p.text.strip())
    body_paras = [p for p in paras if _heading_level(p) is None and p.text.strip()]

    return DocumentMetadata(
        file_name=file_path.name,
        file_size_bytes=file_path.stat().st_size,
        word_count=_count_words(all_text),
        char_count=len(all_text.replace(" ", "").replace("\n", "")),
        paragraph_count=len(body_paras),
        heading_count=len(headings),
        image_count=len(images),
        table_count=len(tables),
        reference_count=len(references),
        has_abstract_cn=abstract_cn is not None,
        has_abstract_en=abstract_en is not None,
        has_keywords_cn=bool(keywords_cn),
        has_keywords_en=bool(keywords_en),
    )


# ── 主入口 ─────────────────────────────────────────────────────

def parse_document(file_path: str | Path) -> DocumentModel:
    """
    解析 DOCX 文件，返回 DocumentModel。

    各字段由独立函数解析，任一字段失败不影响其他字段。
    调用方只需要这一个函数，不需要了解内部细节。
    """
    file_path = Path(file_path)
    doc   = DocxDocument(str(file_path))
    paras = doc.paragraphs

    headings    = parse_headings(paras)
    paragraphs  = parse_paragraphs(paras)
    abstract_cn = parse_abstract_cn(paras)
    abstract_en = parse_abstract_en(paras)
    keywords_cn = parse_keywords_cn(paras)
    keywords_en = parse_keywords_en(paras)
    references  = parse_references(paras)
    images      = parse_images(doc, paras)
    tables      = parse_tables(doc)
    title       = parse_title(paras, headings)
    metadata    = parse_metadata(
        file_path, paras, headings, images, tables,
        references, abstract_cn, abstract_en, keywords_cn, keywords_en,
    )

    return DocumentModel(
        title=title,
        abstract_cn=abstract_cn,
        abstract_en=abstract_en,
        keywords_cn=keywords_cn,
        keywords_en=keywords_en,
        headings=headings,
        paragraphs=paragraphs,
        images=images,
        tables=tables,
        references=references,
        metadata=metadata,
    )
